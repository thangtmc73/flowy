import os
import re
import json
import logging
import time
import base64
import asyncio
import contextvars
import hashlib
import ipaddress
import socket
from datetime import datetime
from difflib import SequenceMatcher
from html.parser import HTMLParser
from io import BytesIO
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
from urllib.parse import urlparse
from urllib.request import Request, urlopen

from dotenv import load_dotenv
from greennode_agentbase.core.logging import configure_logger
from langchain_openai import ChatOpenAI
from langchain.agents import create_agent
from langchain_core.tools import tool

from greennode_agentbase import (
    GreenNodeAgentBaseApp,
    RequestContext,
    PingStatus,
)
from greennode_agentbase.memory import MemoryClient
from greennode_agentbase.memory.models import MemoryRecordSearchRequest, MemoryRecordInsertDirectlyRequest
from greennode_agent_bridge import AgentBaseMemoryEvents
from langgraph.config import get_config

load_dotenv()

logger = logging.getLogger("faq_agent")
configure_logger(logger)

app = GreenNodeAgentBaseApp()

# --- FAQ Data Loading (Multi-Partner Support) ---
FAQ_DATA_PATH = os.environ.get("FAQ_DATA_PATH", "knowledge")
FAQ_ALL_ENTRIES = []
FAQ_INDEX = {}

def load_knowledge_base():
    """Load multi-partner FAQ structure from knowledge base."""
    global FAQ_ALL_ENTRIES, FAQ_INDEX
    
    from pathlib import Path
    
    knowledge_dir = Path(FAQ_DATA_PATH)
    index_path = knowledge_dir / "_index.json"
    
    # Fallback: try old single-file format
    if not index_path.exists():
        old_faq_path = knowledge_dir / "faq.json"
        if old_faq_path.exists():
            print(f"Warning: Using old FAQ format from {old_faq_path}. Consider migrating with scripts/migrate_faq.py")
            load_legacy_format(old_faq_path)
            return
        else:
            print(f"Error: Neither {index_path} nor {old_faq_path} found")
            FAQ_ALL_ENTRIES = []
            return
    
    try:
        with open(index_path, "r", encoding="utf-8") as f:
            FAQ_INDEX = json.load(f)
        
        # Load each product file
        for partner in FAQ_INDEX.get("partners", []):
            if not partner.get("active", True):
                continue
            
            for product in partner.get("products", []):
                product_file = knowledge_dir / product["file"]
                if not product_file.exists():
                    print(f"Warning: Product file not found: {product_file}")
                    continue
                
                with open(product_file, "r", encoding="utf-8") as f:
                    product_data = json.load(f)
                
                # Flatten user_questions into searchable entries
                for faq in product_data.get("faqs", []):
                    for user_question in faq.get("user_questions", []):
                        FAQ_ALL_ENTRIES.append({
                            "question": user_question,
                            "canonical": faq.get("canonical_question", user_question),
                            "answer": faq.get("answer", ""),
                            "category": faq.get("category", "General"),
                            "partner": partner.get("partner_name", "Unknown"),
                            "partner_id": partner.get("partner_id", ""),
                            "product": product.get("product_name", "Unknown"),
                            "product_id": product.get("product_id", ""),
                            "tags": faq.get("tags", []),
                            "priority": faq.get("priority", 5),
                            "id": faq.get("id", "")
                        })
        
        # Load cross-product FAQs if exists
        cross_dir = knowledge_dir / "cross_product"
        if cross_dir.exists():
            for cross_file in cross_dir.glob("*.json"):
                with open(cross_file, "r", encoding="utf-8") as f:
                    cross_data = json.load(f)
                    for faq in cross_data.get("faqs", []):
                        for user_question in faq.get("user_questions", []):
                            FAQ_ALL_ENTRIES.append({
                                "question": user_question,
                                "canonical": faq.get("canonical_question", user_question),
                                "answer": faq.get("answer", ""),
                                "category": faq.get("category", "So sánh / Chung"),
                                "partner": "Zalopay",
                                "product": "Cross-product",
                                "tags": faq.get("tags", []),
                                "priority": faq.get("priority", 8),
                                "id": faq.get("id", "")
                            })
        
        partner_count = len([p for p in FAQ_INDEX.get("partners", []) if p.get("active", True)])
        print(f"✓ Loaded {len(FAQ_ALL_ENTRIES)} FAQ entries from {partner_count} partner(s)")
    
    except Exception as e:
        print(f"Error loading knowledge base: {e}")
        FAQ_ALL_ENTRIES = []


def load_legacy_format(faq_path):
    """Fallback loader for old single-file format."""
    global FAQ_ALL_ENTRIES
    try:
        with open(faq_path, "r", encoding="utf-8") as f:
            old_data = json.load(f)
        
        for item in old_data:
            FAQ_ALL_ENTRIES.append({
                "question": item.get("instruction", ""),
                "canonical": item.get("instruction", ""),
                "answer": item.get("output", ""),
                "category": item.get("category", "General"),
                "partner": "Legacy Partner",
                "product": "Legacy Product",
                "priority": 5
            })
        
        print(f"✓ Loaded {len(FAQ_ALL_ENTRIES)} FAQ entries (legacy format)")
    except Exception as e:
        print(f"Error loading legacy FAQ: {e}")
        FAQ_ALL_ENTRIES = []

def similarity_score(str1: str, str2: str) -> float:
    """Calculate similarity between two strings (0-1)"""
    return SequenceMatcher(None, str1.lower(), str2.lower()).ratio()

def search_faq_fuzzy(query, threshold=0.4, top_k=3, partner_filter=None, category_filter=None):
    """
    Search FAQ using fuzzy matching with multi-partner support.
    
    Args:
        query: User question
        threshold: Minimum similarity score (0-1)
        top_k: Number of results to return
        partner_filter: Filter by partner_id (e.g., "msig")
        category_filter: Filter by category (e.g., "health")
    
    Returns:
        List of matching FAQ entries with scores
    """
    results = []
    
    for entry in FAQ_ALL_ENTRIES:
        # Apply filters
        if partner_filter and entry.get("partner_id") != partner_filter:
            continue
        if category_filter and entry.get("product_category") != category_filter:
            continue
        
        score = similarity_score(query, entry["question"])
        if score >= threshold:
            results.append({
                **entry,
                "score": score
            })
    
    # Rank by similarity first; priority only breaks ties
    results.sort(key=lambda x: (x["score"], x.get("priority", 5)), reverse=True)
    return results[:top_k]


FAQ_SEARCH_TOP_K = 2
FAQ_SEARCH_THRESHOLD = 0.5
FAQ_ANSWER_MAX_CHARS = 600
FAQ_LIST_ANSWER_MAX_CHARS = 800
CATALOG_FAQ_IDS = {"general_products_001", "compare_003", "compare_002"}
PRODUCT_LIST_KEYWORDS = (
    "những gói",
    "các gói",
    "gói nào",
    "gói gì",
    "liệt kê",
    "danh sách",
    "có những",
    "phân phối những",
    "sản phẩm bảo hiểm nào",
    "loại bảo hiểm nào",
)


def _is_product_list_query(query: str) -> bool:
    q = query.lower()
    return "bảo hiểm" in q and any(kw in q for kw in PRODUCT_LIST_KEYWORDS)


def _find_catalog_faq_results(query: str) -> List[Dict[str, Any]]:
    """Return catalog FAQ entries for product-listing questions."""
    best_by_id: Dict[str, Dict[str, Any]] = {}
    for entry in FAQ_ALL_ENTRIES:
        if entry.get("id") not in CATALOG_FAQ_IDS:
            continue
        score = similarity_score(query, entry["question"])
        faq_id = entry["id"]
        if faq_id not in best_by_id or score > best_by_id[faq_id]["score"]:
            best_by_id[faq_id] = {**entry, "score": score}

    if not best_by_id:
        return []

    results = list(best_by_id.values())
    results.sort(key=lambda x: (x["score"], x.get("priority", 5)), reverse=True)
    return results

ZALOPAY_SUPPORT_CONTACT = (
    "**Liên hệ hỗ trợ Zalopay:**\n"
    "* **Tổng đài:** 1900 54 54 36 (cước phí 1.000 VNĐ/phút, hoạt động 24/7 kể cả ngày lễ, Tết)\n"
    "* **Email:** hotro@zalopay.vn\n"
    "* **Trên app Zalopay:** Cá nhân → Trung tâm hỗ trợ → Gửi yêu cầu hỗ trợ\n"
    "* **Facebook:** https://www.facebook.com/Zalopay\n"
    "* **Tổng đài gọi ra (khi Zalopay liên hệ lại):** 028 7307 7068 / 028 7307 5068\n\n"
    "*Lưu ý:* Tổng đài Zalopay hỗ trợ vấn đề về app, thanh toán và dịch vụ trên nền tảng. "
    "Với **bồi thường bảo hiểm**, hãy liên hệ hotline của công ty bảo hiểm đối tác (MSIG, VBI, GIC, Bảo Việt...)."
)

def normalize_brand_name(text: str) -> str:
    """Ensure the platform brand is spelled Zalopay in agent-facing text."""
    if not text:
        return text
    text = re.sub(r"\bZaloPay\b", "Zalopay", text)
    text = re.sub(r"\bZalo Pay\b", "Zalopay", text)
    text = re.sub(r"\bZALOPAY\b", "Zalopay", text)
    return text


def _truncate_faq_answer(answer: str, max_chars: int = FAQ_ANSWER_MAX_CHARS) -> str:
    """Trim long FAQ answers so the LLM context stays focused."""
    if len(answer) <= max_chars:
        return answer
    return answer[:max_chars].rstrip() + "..."


# --- URL Processing Functions ---

URL_PATTERN = re.compile(r"https?://[^\s<>\"')\]]+", re.IGNORECASE)
URL_FETCH_TIMEOUT_SEC = 15
URL_MAX_BYTES = 512_000
URL_MAX_TEXT_CHARS = 12_000
EXTRACTION_MAX_CHARS = int(os.environ.get("EXTRACTION_MAX_CHARS", "10000"))
EXTRACTION_CACHE_MAX_ENTRIES = int(os.environ.get("EXTRACTION_CACHE_MAX_ENTRIES", "128"))

INSURANCE_INFO_FIELDS = (
    "product_name",
    "partner",
    "category",
    "coverage",
    "cost",
    "age_range",
    "benefits",
    "exclusions",
    "highlights",
)

_JSON_FIELD_ALIASES: Dict[str, Tuple[str, ...]] = {
    "product_name": ("product_name", "ten_san_pham", "san_pham", "product", "name", "title"),
    "partner": ("partner", "partner_name", "insurer", "company", "doi_tac", "cong_ty", "provider"),
    "category": ("category", "loai", "type", "product_type", "insurance_type"),
    "coverage": ("coverage", "quyen_loi", "quyền_lợi", "main_coverage", "sum_insured"),
    "cost": ("cost", "price", "premium", "phi", "phí", "fee"),
    "age_range": ("age_range", "age", "do_tuoi", "độ_tuổi"),
    "benefits": ("benefits", "benefit", "quyen_loi_chi_tiet", "benefit_list"),
    "exclusions": ("exclusions", "exclusion", "loai_tru", "loại_trừ"),
    "highlights": ("highlights", "highlight", "diem_noi_bat", "điểm_nổi_bật", "features"),
}

_INSURANCE_KEYWORDS = (
    "bảo hiểm",
    "quyền lợi",
    "phí",
    "premium",
    "coverage",
    "insurance",
    "điều kiện",
    "loại trừ",
    "mức trách nhiệm",
    "hợp đồng",
    "thụ hưởng",
)

_extraction_cache: Dict[str, Dict[str, Any]] = {}

_LATEX_INLINE_RE = re.compile(r"\$([^$\n]+)\$")
_LATEX_DISPLAY_RE = re.compile(r"\$\$(.+?)\$\$|\\\[(.+?)\\\]", re.DOTALL)
_LATEX_PAREN_RE = re.compile(r"\\\((.+?)\\\)", re.DOTALL)
_LATEX_ENV_RE = re.compile(r"\\begin\{[^}]+\}(.*?)\\end\{[^}]+\}", re.DOTALL)
_LATEX_CMD_RE = re.compile(r"\\[a-zA-Z]+\*?(?:\{[^{}]*\})*")
_LATEX_SIMPLE_REPLACEMENTS = (
    (re.compile(r"\\rightarrow|\\to\b"), "→"),
    (re.compile(r"\\leq|\\le\b"), "≤"),
    (re.compile(r"\\geq|\\ge\b"), "≥"),
    (re.compile(r"\\times\b"), "×"),
    (re.compile(r"\\text\{([^{}]*)\}"), r"\1"),
    (re.compile(r"\\textbf\{([^{}]*)\}"), r"\1"),
    (re.compile(r"\\emph\{([^{}]*)\}"), r"\1"),
    (re.compile(r"\\frac\{([^{}]*)\}\{([^{}]*)\}"), r"\1/\2"),
)


def _clean_latex_commands(text: str) -> str:
    """Convert or drop LaTeX commands while keeping plain text/numbers."""
    cleaned = text
    for pattern, replacement in _LATEX_SIMPLE_REPLACEMENTS:
        cleaned = pattern.sub(replacement, cleaned)
    return _LATEX_CMD_RE.sub(" ", cleaned)


def _strip_latex_artifacts(text: str) -> str:
    """Unwrap LaTeX/math delimiters from PDF and HTML sources, keep inner content."""
    if not text:
        return text

    def unwrap_inline(match: re.Match) -> str:
        return _clean_latex_commands(match.group(1).strip())

    def unwrap_display(match: re.Match) -> str:
        content = (match.group(1) or match.group(2) or "").strip()
        return _clean_latex_commands(content)

    def unwrap_env(match: re.Match) -> str:
        return _clean_latex_commands(match.group(1).strip())

    cleaned = _LATEX_DISPLAY_RE.sub(unwrap_display, text)
    cleaned = _LATEX_INLINE_RE.sub(unwrap_inline, cleaned)
    cleaned = _LATEX_PAREN_RE.sub(unwrap_inline, cleaned)
    cleaned = _LATEX_ENV_RE.sub(unwrap_env, cleaned)
    cleaned = _clean_latex_commands(cleaned)
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


class _HTMLTextExtractor(HTMLParser):
    """Strip HTML tags and keep readable text."""

    _SKIP_TAGS = frozenset(
        ("script", "style", "noscript", "math", "svg", "annotation", "semantics")
    )

    def __init__(self):
        super().__init__()
        self._parts: List[str] = []
        self._skip_stack: List[str] = []

    def handle_starttag(self, tag, attrs):
        if tag in self._SKIP_TAGS:
            self._skip_stack.append(tag)
            return

        attr_map = {name.lower(): value for name, value in attrs}
        class_value = (attr_map.get("class") or "").lower()
        if any(token in class_value for token in ("katex", "mathjax", "math-inline", "math-display")):
            self._skip_stack.append("math")
            return

        if tag in ("p", "div", "br", "li", "h1", "h2", "h3", "h4", "tr", "section"):
            self._parts.append("\n")

    def handle_endtag(self, tag):
        if not self._skip_stack:
            return
        top = self._skip_stack[-1]
        if (tag in self._SKIP_TAGS and top == tag) or (
            tag in ("span", "div") and top == "math"
        ):
            self._skip_stack.pop()

    def handle_data(self, data):
        if not self._skip_stack and data.strip():
            self._parts.append(data)

    def get_text(self) -> str:
        text = "".join(self._parts)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()


def extract_urls_from_message(message: str) -> List[str]:
    """Return unique HTTP(S) URLs found in a user message."""
    seen = set()
    urls = []
    for match in URL_PATTERN.findall(message or ""):
        url = match.rstrip(".,;:!?)")
        if url not in seen:
            seen.add(url)
            urls.append(url)
    return urls


def _is_private_ip(ip_str: str) -> bool:
    try:
        ip = ipaddress.ip_address(ip_str)
    except ValueError:
        return True
    return (
        ip.is_private
        or ip.is_loopback
        or ip.is_link_local
        or ip.is_reserved
        or ip.is_multicast
    )


def _validate_public_http_url(url: str) -> str:
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise ValueError("Chỉ hỗ trợ link http hoặc https.")

    hostname = parsed.hostname
    if not hostname:
        raise ValueError("Link không hợp lệ.")

    if hostname.lower() in ("localhost", "127.0.0.1", "0.0.0.0", "::1"):
        raise ValueError("Không thể truy cập link nội bộ.")

    try:
        addr_info = socket.getaddrinfo(hostname, None)
    except socket.gaierror as exc:
        raise ValueError(f"Không thể phân giải tên miền: {hostname}") from exc

    for info in addr_info:
        ip_str = info[4][0]
        if _is_private_ip(ip_str):
            raise ValueError("Không thể truy cập link nội bộ hoặc mạng riêng.")

    return url


def _html_to_text(html: str) -> str:
    parser = _HTMLTextExtractor()
    parser.feed(html)
    return _strip_latex_artifacts(parser.get_text())


def _decode_response_body(raw: bytes, content_type: str) -> str:
    charset = "utf-8"
    if "charset=" in content_type.lower():
        charset = content_type.lower().split("charset=", 1)[1].split(";", 1)[0].strip() or "utf-8"

    try:
        text = raw.decode(charset, errors="replace")
    except LookupError:
        text = raw.decode("utf-8", errors="replace")

    content_type = content_type.lower()
    if "html" in content_type:
        return _html_to_text(text)
    if "json" in content_type:
        try:
            return json.dumps(json.loads(text), ensure_ascii=False, indent=2)
        except json.JSONDecodeError:
            return text
    return text


def fetch_url_content(url: str) -> str:
    """Fetch and normalize text content from a public HTTP(S) URL."""
    safe_url = _validate_public_http_url(url)
    request = Request(
        safe_url,
        headers={
            "User-Agent": "FlowyInsuranceBot/1.0 (+https://zalopay.vn)",
            "Accept": "text/html,application/json,text/plain,*/*",
        },
    )

    with urlopen(request, timeout=URL_FETCH_TIMEOUT_SEC) as response:
        content_type = response.headers.get("Content-Type", "")
        raw = response.read(URL_MAX_BYTES + 1)

    if len(raw) > URL_MAX_BYTES:
        raise ValueError("Nội dung link quá lớn để xử lý.")

    text = _decode_response_body(raw, content_type)
    if not text.strip():
        raise ValueError("Không đọc được nội dung hữu ích từ link.")

    return text[:URL_MAX_TEXT_CHARS]


async def fetch_url_content_async(url: str) -> str:
    return await asyncio.to_thread(fetch_url_content, url)


def build_insurance_comparison_message(
    source_type: str,
    source_name: str,
    insurance_info: Dict[str, Any],
    user_message: str,
) -> str:
    """Build a prompt that asks the agent to compare extracted insurance info."""
    return f"""Người dùng đã cung cấp {source_type}: {source_name}

Thông tin trích xuất:
{json.dumps(insurance_info, ensure_ascii=False, indent=2)}

Câu hỏi của người dùng: {user_message}

Hãy phân tích nội dung này và so sánh với các sản phẩm bảo hiểm trên Zalopay.
Sử dụng tool 'compare_insurance_products' để đưa ra so sánh chi tiết."""


# --- File Processing Functions ---

def _normalize_insurance_info(data: Dict[str, Any]) -> Dict[str, Any]:
    """Ensure extracted insurance info has the expected schema."""
    normalized = {field: data.get(field) for field in INSURANCE_INFO_FIELDS}
    for field, value in normalized.items():
        if value is None:
            normalized[field] = "N/A" if field in ("exclusions", "highlights") else ""
        elif isinstance(value, (list, dict)):
            normalized[field] = json.dumps(value, ensure_ascii=False)
        else:
            normalized[field] = _strip_latex_artifacts(str(value).strip())
    if not normalized["product_name"]:
        normalized["product_name"] = "Unknown"
    return normalized


def _pick_json_value(data: Dict[str, Any], aliases: Tuple[str, ...]) -> Any:
    """Find the first matching value from common JSON field names."""
    lowered = {str(k).lower(): v for k, v in data.items()}
    for alias in aliases:
        if alias in data and data[alias] not in (None, ""):
            return data[alias]
        if alias.lower() in lowered and lowered[alias.lower()] not in (None, ""):
            return lowered[alias.lower()]
    return None


def extract_insurance_from_json(data: Dict[str, Any]) -> Dict[str, Any]:
    """Map structured JSON upload/API response to insurance info without LLM."""
    if all(field in data for field in INSURANCE_INFO_FIELDS):
        return _normalize_insurance_info(data)

    mapped = {
        field: _pick_json_value(data, aliases)
        for field, aliases in _JSON_FIELD_ALIASES.items()
    }

    # Nested product_summary blocks (knowledge-base style)
    summary = data.get("product_summary")
    if isinstance(summary, dict):
        for field, aliases in _JSON_FIELD_ALIASES.items():
            if not mapped.get(field):
                mapped[field] = _pick_json_value(summary, aliases)

    return _normalize_insurance_info(mapped)


def _extract_pdf_text(content: Dict[str, Any]) -> str:
    """Extract plain text from a base64-encoded PDF payload."""
    from pypdf import PdfReader

    pdf_b64 = content.get("data", "")
    if not pdf_b64:
        raise ValueError("PDF không có dữ liệu.")

    try:
        raw = base64.b64decode(pdf_b64)
    except (ValueError, TypeError) as exc:
        raise ValueError("PDF không hợp lệ.") from exc

    reader = PdfReader(BytesIO(raw))
    parts = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            parts.append(page_text.strip())

    text = "\n\n".join(parts).strip()
    if not text:
        raise ValueError(
            "Không trích xuất được nội dung text từ PDF (có thể là file scan/ảnh)."
        )
    return _strip_latex_artifacts(text)


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
LEGACY_DOC_MIME = "application/msword"


def _is_pdf_upload(file_type: str, content: Any) -> bool:
    return file_type == "application/pdf" or (
        isinstance(content, dict) and content.get("type") == "pdf"
    )


def _is_docx_upload(file_type: str, content: Any, file_name: str = "") -> bool:
    if file_type == DOCX_MIME:
        return True
    if isinstance(content, dict) and content.get("type") == "docx":
        return True
    return file_name.lower().endswith(".docx")


def _is_legacy_doc_upload(file_type: str, file_name: str = "") -> bool:
    lowered_name = file_name.lower()
    if lowered_name.endswith(".docx"):
        return False
    return file_type == LEGACY_DOC_MIME or lowered_name.endswith(".doc")


def _is_json_upload(file_type: str, content: Any) -> bool:
    if file_type == "application/json":
        return True
    if not isinstance(content, dict):
        return False
    return content.get("type") not in ("pdf", "docx")


def _extract_docx_text(content: Dict[str, Any]) -> str:
    """Extract plain text from a base64-encoded DOCX payload."""
    from docx import Document

    docx_b64 = content.get("data", "")
    if not docx_b64:
        raise ValueError("DOCX không có dữ liệu.")

    try:
        raw = base64.b64decode(docx_b64)
    except (ValueError, TypeError) as exc:
        raise ValueError("DOCX không hợp lệ.") from exc

    document = Document(BytesIO(raw))
    parts: List[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n\n".join(parts).strip()
    if not text:
        raise ValueError("Không trích xuất được nội dung text từ DOCX.")
    return _strip_latex_artifacts(text)


def parse_uploaded_file(file_data: Dict[str, Any]) -> str:
    """Parse uploaded file content based on file type."""
    file_type = file_data.get("type", "")
    file_name = file_data.get("name", "")
    content = file_data.get("content", "")

    if _is_legacy_doc_upload(file_type, file_name):
        raise ValueError(
            "Chỉ hỗ trợ file Word .docx. Vui lòng lưu lại file dạng .docx."
        )

    if _is_json_upload(file_type, content):
        if isinstance(content, dict):
            return json.dumps(content, ensure_ascii=False, indent=2)
        return str(content)

    if _is_pdf_upload(file_type, content):
        pdf_payload = content if isinstance(content, dict) else {"data": content}
        return _extract_pdf_text(pdf_payload)

    if _is_docx_upload(file_type, content, file_name):
        docx_payload = content if isinstance(content, dict) else {"data": content}
        return _extract_docx_text(docx_payload)

    if isinstance(content, str):
        return content

    return str(content)


def _select_text_for_extraction(text: str, max_chars: int) -> str:
    """Prefer paragraphs with insurance-related keywords over raw leading text."""
    text = _strip_latex_artifacts(text.strip())
    if len(text) <= max_chars:
        return text

    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    if not paragraphs:
        return text[:max_chars]

    scored: List[Tuple[int, int, str]] = []
    for index, paragraph in enumerate(paragraphs):
        lowered = paragraph.lower()
        score = sum(1 for keyword in _INSURANCE_KEYWORDS if keyword in lowered)
        scored.append((score, index, paragraph))

    scored.sort(key=lambda item: (-item[0], item[1]))

    selected: List[Tuple[int, str]] = []
    total_len = 0
    for score, index, paragraph in scored:
        if score == 0 and selected:
            continue
        addition = len(paragraph) + (2 if selected else 0)
        if total_len + addition > max_chars:
            if not selected:
                return text[:max_chars]
            break
        selected.append((index, paragraph))
        total_len += addition

    if not selected:
        return text[:max_chars]

    selected.sort(key=lambda item: item[0])
    combined = "\n\n".join(paragraph for _, paragraph in selected)
    if len(combined) > max_chars:
        return combined[:max_chars]
    return combined


def _parse_llm_json_content(content: str) -> Dict[str, Any]:
    """Parse JSON from LLM response, stripping markdown fences if present."""
    content = content.strip()
    if "```json" in content:
        content = content.split("```json", 1)[1].split("```", 1)[0].strip()
    elif "```" in content:
        content = content.split("```", 1)[1].split("```", 1)[0].strip()
    parsed = json.loads(content)
    if not isinstance(parsed, dict):
        raise ValueError("LLM response is not a JSON object")
    return parsed


def _build_extraction_cache_key(user_id: str, session_id: str, source_key: str) -> str:
    digest = hashlib.sha256(source_key.encode("utf-8")).hexdigest()[:16]
    return f"{user_id}:{session_id}:{digest}"


def _get_cached_extraction(cache_key: Optional[str]) -> Optional[Dict[str, Any]]:
    if not cache_key:
        return None
    return _extraction_cache.get(cache_key)


def _set_cached_extraction(cache_key: Optional[str], data: Dict[str, Any]) -> None:
    if not cache_key:
        return
    if len(_extraction_cache) >= EXTRACTION_CACHE_MAX_ENTRIES:
        oldest_key = next(iter(_extraction_cache))
        _extraction_cache.pop(oldest_key, None)
    _extraction_cache[cache_key] = data


def extract_insurance_features(
    text_content: str,
    llm: ChatOpenAI,
    cache_key: Optional[str] = None,
) -> Dict[str, Any]:
    """Use analysis LLM to extract insurance features from unstructured text."""
    cached = _get_cached_extraction(cache_key)
    if cached is not None:
        logger.info("[extract] cache hit key=%s", cache_key)
        return cached

    selected_text = _select_text_for_extraction(text_content, EXTRACTION_MAX_CHARS)
    extraction_prompt = f"""Analyze this insurance document and extract key information in JSON format.

Document content:
{selected_text}

Notes:
- The source may contain LaTeX/math markup ($...$, \\rightarrow, etc.) from PDF or HTML rendering.
- Ignore LaTeX syntax; extract plain Vietnamese meaning only. Do not copy LaTeX into JSON values.

Extract the following information (in Vietnamese):
- product_name: Tên sản phẩm bảo hiểm
- partner: Công ty bảo hiểm/đối tác
- category: Loại bảo hiểm (sức khỏe, xe, du lịch, cyber, v.v.)
- coverage: Quyền lợi bảo hiểm chính
- cost: Chi phí/phí bảo hiểm
- age_range: Độ tuổi áp dụng
- benefits: Các quyền lợi cụ thể
- exclusions: Điều khoản loại trừ (nếu có)
- highlights: Điểm nổi bật

Return ONLY valid JSON, no additional text."""

    try:
        response = llm.invoke(extraction_prompt)
        parsed = _parse_llm_json_content(response.content)
        result = _normalize_insurance_info(parsed)
        _set_cached_extraction(cache_key, result)
        return result
    except Exception as e:
        logger.error("Error extracting features with LLM: %s", e)
        return {
            "product_name": "Unknown",
            "error": str(e),
            "raw_content": text_content[:500],
        }


def resolve_insurance_info(
    *,
    llm: ChatOpenAI,
    text_content: Optional[str] = None,
    json_data: Optional[Dict[str, Any]] = None,
    cache_key: Optional[str] = None,
) -> Dict[str, Any]:
    """Resolve insurance info from JSON (no LLM) or unstructured text (analysis LLM)."""
    if json_data is not None:
        logger.info("[extract] using structured JSON mapping (skip analysis LLM)")
        return extract_insurance_from_json(json_data)

    if not text_content or not text_content.strip():
        raise ValueError("Không có nội dung để phân tích.")

    try:
        parsed = json.loads(text_content)
        if isinstance(parsed, dict):
            logger.info("[extract] detected JSON text payload (skip analysis LLM)")
            return extract_insurance_from_json(parsed)
    except json.JSONDecodeError:
        pass

    logger.info(
        "[extract] using analysis LLM chars=%d (selected from %d)",
        min(len(text_content), EXTRACTION_MAX_CHARS),
        len(text_content),
    )
    return extract_insurance_features(text_content, llm, cache_key=cache_key)


def load_all_products() -> List[Dict[str, Any]]:
    """Load all insurance products from knowledge base.

    Returns:
        List of product dicts with partner info
    """
    products = []
    knowledge_dir = Path(FAQ_DATA_PATH)
    index_path = knowledge_dir / "_index.json"
    
    if not index_path.exists():
        return products
    
    try:
        with open(index_path, "r", encoding="utf-8") as f:
            index_data = json.load(f)
        
        for partner in index_data.get("partners", []):
            if not partner.get("active", True):
                continue
            
            for product in partner.get("products", []):
                product_file = knowledge_dir / product["file"]
                if not product_file.exists():
                    continue
                
                with open(product_file, "r", encoding="utf-8") as f:
                    product_data = json.load(f)
                
                products.append({
                    "partner_id": partner.get("partner_id", ""),
                    "partner_name": partner.get("partner_name", ""),
                    "product_id": product.get("product_id", ""),
                    "product_name": product.get("product_name", ""),
                    "category": product.get("category", ""),
                    "data": product_data
                })
        
        return products
    except Exception as e:
        logger.error(f"Error loading products: {e}")
        return products


# Load FAQ data on startup
load_knowledge_base()

# --- Memory Configuration ---
# Create a memory with: /agentbase-memory
# Set the memory ID here or via MEMORY_ID env var
MEMORY_ID = os.environ.get("MEMORY_ID", "")
if not MEMORY_ID:
    raise ValueError("MEMORY_ID environment variable is required for memory-enabled agents")

# Strategy ID for long-term memory namespace partitioning
# This is fixed per memory instance — do NOT pass as a tool parameter
MEMORY_STRATEGY_ID = os.environ.get("MEMORY_STRATEGY_ID", "default")

# CheckpointSaver: persists conversation state as events in AgentBase Memory
# This enables multi-turn conversations that survive restarts
checkpointer = AgentBaseMemoryEvents(memory_id=MEMORY_ID)

# MemoryClient: used by long-term memory tools to store/search semantic facts
memory_client = MemoryClient()

# --- LLM Configuration ---
# Uses any OpenAI-compatible LLM provider (GreenNode AIP, OpenAI, Ollama, etc.)
# Set LLM_BASE_URL, LLM_API_KEY, and LLM_MODEL in your .env file.
# For GreenNode AIP: use /agentbase-llm to manage API keys and browse models.
# For other providers: set the appropriate base URL and API key.
# Production: use /agentbase-identity to store API key, inject via @requires_api_key
#
# Dual-model routing (same API key / base URL):
# - LLM_MODEL: chat agent (FAQ replies, comparison answers) — default Gemma for low latency
# - LLM_MODEL_ANALYSIS: file upload / external link extraction — default MiniMax for document parsing
LLM_MODEL = os.environ.get("LLM_MODEL", "")
LLM_MODEL_ANALYSIS = os.environ.get("LLM_MODEL_ANALYSIS", "minimax/minimax-m2.5")
LLM_BASE_URL = os.environ.get("LLM_BASE_URL", "")
LLM_API_KEY = os.environ.get("LLM_API_KEY", "")
if not LLM_MODEL or not LLM_BASE_URL or not LLM_API_KEY:
    raise ValueError(
        "LLM_MODEL, LLM_BASE_URL, and LLM_API_KEY environment variables are required. "
        "Set them in your .env file or use /agentbase-llm to get a platform API key."
    )


LLM_REQUEST_TIMEOUT_SEC = float(os.environ.get("LLM_REQUEST_TIMEOUT_SEC", "90"))
AGENT_INVOKE_TIMEOUT_SEC = float(os.environ.get("AGENT_INVOKE_TIMEOUT_SEC", "110"))
AGENT_INVOKE_WARN_INTERVAL_SEC = float(os.environ.get("AGENT_INVOKE_WARN_INTERVAL_SEC", "30"))

_invoke_last_tool: contextvars.ContextVar[Optional[str]] = contextvars.ContextVar(
    "invoke_last_tool", default=None
)


def _mark_tool_started(tool_name: str) -> None:
    _invoke_last_tool.set(tool_name)


def _create_llm(model: str) -> ChatOpenAI:
    return ChatOpenAI(
        model=model,
        base_url=LLM_BASE_URL,
        api_key=LLM_API_KEY,
        timeout=LLM_REQUEST_TIMEOUT_SEC,
    )


llm_chat = _create_llm(LLM_MODEL)
llm_analysis = _create_llm(LLM_MODEL_ANALYSIS)
logger.info(
    "[llm] chat=%s analysis=%s base_url=%s",
    LLM_MODEL,
    LLM_MODEL_ANALYSIS,
    LLM_BASE_URL,
)


# --- Long-Term Memory Tools (via MemoryClient SDK) ---
# actor_id: retrieved from LangGraph configurable (set in handler via context.user_id)
# strategy_id: app-level config (MEMORY_STRATEGY_ID), fixed per memory instance
# Neither should be exposed as tool parameters to avoid LLM hallucination


def _get_actor_id():
    """Get actor_id from LangGraph configurable (set during graph.invoke)."""
    try:
        config = get_config()
        return config["configurable"].get("actor_id", "default")
    except RuntimeError:
        # Called outside runnable context - return default
        return "default"


def _build_namespace(actor_id: str) -> str:
    """Build memory namespace from strategy_id (app config) and actor_id (runtime config)."""
    return f"/strategies/{MEMORY_STRATEGY_ID}/actors/{actor_id}"


@tool
async def remember(fact: str) -> str:
    """Store a fact in long-term memory for later retrieval.

    Args:
        fact: The fact or information to remember.
    """
    _mark_tool_started("remember")
    actor_id = _get_actor_id()
    namespace = _build_namespace(actor_id)
    logger.info("[tool:remember] actor=%s fact=%r", actor_id, fact[:80])
    await memory_client.insert_memory_records_directly_async(
        id=MEMORY_ID,
        namespace=namespace,
        request=MemoryRecordInsertDirectlyRequest(memory_records=[fact]),
    )
    logger.info("[tool:remember] stored OK")
    return f"Remembered: {fact}"


@tool
async def recall(query: str) -> str:
    """Search long-term memory for facts relevant to a query.

    Args:
        query: Natural language search query.
    """
    _mark_tool_started("recall")
    actor_id = _get_actor_id()
    namespace = _build_namespace(actor_id)
    logger.info("[tool:recall] actor=%s query=%r", actor_id, query[:80])
    t0 = time.monotonic()
    results = await memory_client.search_memory_records_async(
        id=MEMORY_ID,
        namespace=namespace,
        request=MemoryRecordSearchRequest(query=query, limit=10),
    )
    logger.info("[tool:recall] found %d results (%.2fs)", len(results), time.monotonic() - t0)
    if not results:
        return "No relevant memories found."
    # SDK 1.0.3 returns dicts; older versions return model objects — handle both
    def _get(r, key, default=""):
        return r[key] if isinstance(r, dict) else getattr(r, key, default)
    return "\n".join(f"- {_get(r, 'memory')} (score: {_get(r, 'score', 0):.2f})" for r in results)


@tool
def search_faq_docs(query: str) -> str:
    """Search FAQ documentation for answers to user questions about insurance products on Zalopay.

    Supports multiple partners and product types (health, car, travel, cyber, etc.).
    Automatically detects partner and product context from the query.

    Args:
        query: The user's question or topic to search for in the FAQ docs.
    """
    _mark_tool_started("search_faq_docs")
    logger.info("[tool:search_faq] query=%r", query[:80])
    if not FAQ_ALL_ENTRIES:
        return f"FAQ database is not loaded. Please check if {FAQ_DATA_PATH} exists."

    # Smart context detection
    partner_filter = None
    query_lower = query.lower()
    
    # Detect specific partner mentions in query
    partner_patterns = {
        "msig": ["msig", "sức khỏe 24/7"],
        "gic": ["gic", "credit topup", "sống tự tin"],
        "vbi": ["vbi", "cyber"],
        "baoviet": ["bảo việt", "baoviet", "chuyến bay", "trễ chuyến", "hủy chuyến", "saladin"],
    }
    
    for partner_id, patterns in partner_patterns.items():
        if any(pattern in query_lower for pattern in patterns):
            partner_filter = partner_id
            break

    is_product_list = _is_product_list_query(query)
    t0 = time.monotonic()
    if is_product_list and not partner_filter:
        results = _find_catalog_faq_results(query)[:FAQ_SEARCH_TOP_K]
        if not results:
            results = search_faq_fuzzy(
                query,
                threshold=FAQ_SEARCH_THRESHOLD,
                top_k=FAQ_SEARCH_TOP_K,
                partner_filter=partner_filter,
            )
    else:
        results = search_faq_fuzzy(
            query,
            threshold=FAQ_SEARCH_THRESHOLD,
            top_k=FAQ_SEARCH_TOP_K,
            partner_filter=partner_filter,
        )
    logger.info("[tool:search_faq] matched %d results (%.2fs)", len(results), time.monotonic() - t0)

    if not results:
        return (
            "Không tìm thấy thông tin phù hợp trong FAQ về sản phẩm bảo hiểm trên Zalopay.\n"
            "Có thể hỏi về:\n"
            "- Các gói bảo hiểm sức khỏe\n"
            "- Bảo hiểm an ninh mạng (Cyber)\n"
            "- Bảo hiểm tài chính (Credit Topup)\n"
            "- Bảo hiểm trễ và hủy chuyến bay (Bảo Việt)\n"
            "- Quyền lợi, chi phí, độ tuổi áp dụng\n"
            "- Quy trình mua và bồi thường\n\n"
            f"{ZALOPAY_SUPPORT_CONTACT}"
        )

    best = results[0]
    confidence_pct = int(best["score"] * 100)
    answer_max_chars = FAQ_LIST_ANSWER_MAX_CHARS if is_product_list else FAQ_ANSWER_MAX_CHARS
    response = (
        f"Kết quả phù hợp nhất ({confidence_pct}%):\n"
        f"Sản phẩm: {best['partner']} - {best['product']}\n"
        f"Câu hỏi: {best['canonical']}\n"
        f"Danh mục: {best['category']}\n"
        f"Trả lời: {_truncate_faq_answer(best['answer'], answer_max_chars)}\n"
    )

    if len(results) > 1:
        response += "\nKết quả liên quan khác (chỉ tham khảo, không trả lời trừ khi user hỏi):\n"
        for result in results[1:]:
            other_pct = int(result["score"] * 100)
            response += (
                f"- {result['partner']} - {result['product']}: "
                f"{result['canonical']} ({result['category']}, {other_pct}%)\n"
            )

    return response


@tool
async def compare_insurance_products(uploaded_info: str) -> str:
    """Compare insurance info from uploaded files or external links with existing products on Zalopay.
    
    This tool analyzes insurance information extracted from user-provided content and compares
    it with products available on the Zalopay platform.
    
    Args:
        uploaded_info: Extracted insurance information in JSON string format
    """
    _mark_tool_started("compare_insurance_products")
    logger.info("[tool:compare_insurance] starting comparison")
    
    try:
        # Parse uploaded info
        if isinstance(uploaded_info, str):
            uploaded = json.loads(uploaded_info)
        else:
            uploaded = uploaded_info
        
        # Load all existing products
        products = load_all_products()
        
        if not products:
            return "Không thể tải danh sách sản phẩm từ knowledge base."
        
        # Build comparison response
        response = f"## So sánh: {uploaded.get('product_name', 'Sản phẩm tải lên')}\n\n"
        response += f"**Đối tác:** {uploaded.get('partner', 'N/A')}\n"
        response += f"**Loại:** {uploaded.get('category', 'N/A')}\n\n"
        
        # Find similar products by category
        similar_products = [
            p for p in products 
            if uploaded.get('category', '').lower() in p.get('category', '').lower() or
               p.get('category', '').lower() in uploaded.get('category', '').lower()
        ]
        
        if not similar_products:
            response += "Không tìm thấy sản phẩm tương tự trong danh mục này trên Zalopay.\n\n"
            response += "**Sản phẩm hiện có trên Zalopay:**\n"
            for p in products[:3]:
                response += f"- {p['product_name']} ({p['partner_name']})\n"
            return response
        
        response += f"**Tìm thấy {len(similar_products)} sản phẩm tương tự trên Zalopay:**\n\n"
        
        # Detailed comparison
        for i, product in enumerate(similar_products[:3], 1):
            product_data = product['data']
            response += f"### {i}. {product['product_name']} - {product['partner_name']}\n\n"
            
            # Extract key info from product
            summary = product_data.get('product_summary', {})
            response += f"**Chi phí:** {summary.get('cost', 'N/A')}\n"
            response += f"**Quyền lợi:** {summary.get('coverage', 'N/A')}\n"
            response += f"**Độ tuổi:** {summary.get('age_range', 'N/A')}\n"
            
            # Find specific FAQs about benefits
            benefit_faqs = [
                faq for faq in product_data.get('faqs', [])
                if 'quyền lợi' in faq.get('canonical_question', '').lower() or
                   'benefits' in faq.get('category', '').lower()
            ]
            
            if benefit_faqs:
                response += f"\n**Chi tiết quyền lợi:**\n{benefit_faqs[0].get('answer', '')[:300]}...\n"
            
            response += "\n---\n\n"
        
        # Comparison summary
        response += "## Kết luận\n\n"
        response += f"Sản phẩm tải lên thuộc danh mục **{uploaded.get('category')}**, "
        response += f"tương tự với các sản phẩm của {', '.join([p['partner_name'] for p in similar_products[:2]])} trên Zalopay.\n\n"
        
        response += "Bạn có thể:\n"
        response += "- So sánh chi tiết quyền lợi và chi phí\n"
        response += "- Tìm hiểu thêm về quy trình mua bảo hiểm trên Zalopay\n"
        response += "- Liên hệ để được tư vấn cụ thể hơn\n"
        
        return response
        
    except Exception as e:
        logger.error(f"[tool:compare_insurance] error: {e}")
        return f"Lỗi khi so sánh sản phẩm: {str(e)}"


# --- Create Agent with Checkpointer ---
# create_agent builds a compiled LangGraph StateGraph with tool-calling support.
# checkpointer: persists conversation state via AgentBase Memory (short-term)
# Long-term memory is handled by remember/recall tools via MemoryClient SDK
agent = create_agent(
    llm_chat,
    tools=[remember, recall, search_faq_docs, compare_insurance_products],
    system_prompt=(
        "Bạn là trợ lý tư vấn bảo hiểm trên nền tảng Zalopay.\n\n"
        "QUY TẮC THƯƠNG HIỆU:\n"
        "- Luôn viết tên nền tảng là **Zalopay** (chữ Z viết hoa, còn lại viết thường)\n"
        "- KHÔNG dùng ZaloPay, Zalo Pay hay ZALOPAY\n\n"
        "NGUYÊN TẮC TRẢ LỜI (ưu tiên cao nhất):\n"
        "- Chỉ trả lời đúng phạm vi câu hỏi. Không lan sang quyền lợi, quy trình, so sánh nếu user không hỏi.\n"
        "- Câu hỏi yes/no hoặc một con số → trả lời trực tiếp trong 1–3 câu trước; chi tiết chỉ thêm khi cần.\n"
        "- Tóm tắt từ FAQ thành 2–4 câu; không copy nguyên văn trừ khi user hỏi chi tiết hoặc cần giữ bảng/HTML.\n"
        "- Dùng kết quả FAQ phù hợp nhất; bỏ qua kết quả phụ trừ khi user hỏi 'tất cả', 'so sánh', 'liệt kê'.\n"
        "- Khi user hỏi liệt kê/có những gói bảo hiểm nào → bắt buộc nêu đủ 4 sản phẩm: Cyber (VBI), Sức khỏe 24/7 (MSIG), Sống tự tin (GIC), trễ/hủy chuyến bay (Bảo Việt).\n"
        "- Tối đa ~150 từ cho câu hỏi đơn giản; kết thúc bằng 1 gợi ý ngắn thay vì liệt kê mọi chủ đề liên quan.\n\n"
        "Vai trò của bạn:\n"
        "- Trả lời câu hỏi về các sản phẩm bảo hiểm từ nhiều đối tác bằng cách tìm kiếm trong FAQ với tool 'search_faq_docs'\n"
        "- So sánh sản phẩm bảo hiểm từ file upload hoặc link ngoài với các gói hiện có trên Zalopay bằng tool 'compare_insurance_products'\n"
        "- Nhớ thông tin quan trọng về người dùng với tool 'remember'\n"
        "- Tra cứu lại các tương tác và thông tin đã học với tool 'recall' khi cần cá nhân hóa\n"
        "- Cung cấp câu trả lời chính xác, ngắn gọn, dựa trên tài liệu FAQ\n"
        "- So sánh khách quan giữa các gói bảo hiểm chỉ khi được hỏi\n"
        "- Nếu không tìm thấy câu trả lời trong tài liệu, hãy nói thẳng thắn\n"
        "- Giao tiếp thân thiện, giữ ngữ cảnh xuyên suốt cuộc trò chuyện\n\n"
        "KHI GỢI Ý LIÊN HỆ TỔNG ĐÀI / CSKH ZALOPAY:\n"
        "- Bắt buộc đưa đầy đủ thông tin liên lạc Zalopay (tổng đài, cước phí, giờ hoạt động, email, hỗ trợ trên app, Facebook, số gọi ra)\n"
        "- Tra FAQ bằng search_faq_docs với query 'liên hệ hỗ trợ Zalopay' nếu cần chi tiết\n"
        "- Không chỉ nói 'liên hệ tổng đài' mà phải kèm số hotline và các kênh khác\n"
        "- Phân biệt: Zalopay CSKH cho app/thanh toán; hotline đối tác BH cho bồi thường\n\n"
        "Quy trình trả lời:\n"
        "1. Tìm kiếm FAQ bằng 'search_faq_docs'\n"
        "2. Chỉ gọi 'recall' khi cần thông tin cá nhân hóa từ các lượt chat trước\n"
        "3. Trả lời súc tích theo nguyên tắc trên; gọi 'remember' khi user chia sẻ nhu cầu/sở thích quan trọng\n\n"
        "XỬ LÝ FILE UPLOAD / LINK NGOÀI:\n"
        "- Khi người dùng upload file hoặc gửi link bảo hiểm, phân tích thông tin trong nội dung đó\n"
        "- Tài liệu PDF/HTML đôi khi còn sót cú pháp LaTeX/math ($...$, \\rightarrow, v.v.) — bỏ qua, không copy vào câu trả lời\n"
        "- Sử dụng 'compare_insurance_products' để so sánh với sản phẩm trên Zalopay\n"
        "- Đưa ra nhận xét khách quan về ưu/nhược điểm\n"
        "- Gợi ý sản phẩm tương tự hoặc tốt hơn trên Zalopay\n\n"
        "QUY TẮC FORMAT OUTPUT (MARKDOWN):\n"
        "- Trả lời bằng Markdown; câu hỏi đơn giản thường chỉ cần 1 đoạn hoặc vài bullet ngắn\n"
        "- Dùng **bold** cho thông tin then chốt\n"
        "- KHÔNG dùng LaTeX/math ($...$, \\rightarrow, v.v.). Khi mô tả luồng bước, dùng ký tự → hoặc dấu gạch ngang\n"
        "- Dùng ###, emoji, bảng CHỈ khi user hỏi so sánh, liệt kê nhiều mục, hoặc yêu cầu chi tiết\n"
        "- Dùng markdown tables khi so sánh:\n"
        "  | Tiêu chí | Sản phẩm A | Sản phẩm B |\n"
        "  |----------|------------|------------|\n"
        "  | Chi phí  | 100K       | 150K       |\n\n"
        "FORMAT CHO CÂU HỎI ĐƠN LẺ (1 sản phẩm):\n"
        "- Tóm tắt 2–4 câu; giữ HTML/bảng FAQ chỉ khi user cần chi tiết giá hoặc quyền lợi đầy đủ\n\n"
        "FORMAT CHO CÂU HỎI TỔNG HỢP (nhiều sản phẩm — chỉ khi user hỏi tất cả/so sánh):\n"
        "Dùng format phân cấp rõ ràng:\n\n"
        "🔹 **TÊN SẢN PHẨM 1** (Đối tác)\n"
        "- **Chi phí:** ...\n"
        "- **Quyền lợi:** ...\n"
        "- **Đặc điểm:** ...\n\n"
        "🔹 **TÊN SẢN PHẨM 2** (Đối tác)\n"
        "- **Chi phí:** ...\n"
        "- **Quyền lợi:** ...\n"
        "- **Đặc điểm:** ...\n\n"
        "FORMAT CHO SO SÁNH SẢN PHẨM:\n"
        "### So sánh: [Tên sản phẩm]\n\n"
        "**Sản phẩm upload:**\n"
        "- Đối tác: ...\n"
        "- Chi phí: ...\n"
        "- Quyền lợi: ...\n\n"
        "**Sản phẩm tương tự trên Zalopay:**\n\n"
        "| Tiêu chí | Sản phẩm Upload | Zalopay - A | Zalopay - B |\n"
        "|----|----|----|----|\n"
        "| Chi phí | ... | ... | ... |\n"
        "| Quyền lợi | ... | ... | ... |\n\n"
        "**Nhận xét:**\n"
        "- Ưu điểm: ...\n"
        "- Nhược điểm: ...\n"
        "- Gợi ý: ...\n\n"
        "Lưu ý:\n"
        "- Luôn trả lời bằng tiếng Việt\n"
        "- Thân thiện, nhiệt tình, chuyên nghiệp\n"
        "- Nếu người dùng chào hỏi, hãy giới thiệu bản thân và dịch vụ\n"
        "- Khi so sánh giữa các đối tác, trình bày khách quan, không thiên vị\n"
        "- KHÔNG bao giờ tự thêm markdown syntax không cần thiết"
    ),
    checkpointer=checkpointer,
)


async def _invoke_agent_with_timeout_logging(
    input_state: dict,
    config: dict,
    *,
    user_id: str,
    session_id: str,
    message_preview: str,
) -> dict:
    """Run agent.ainvoke with periodic slow-path warnings and a hard timeout."""
    _invoke_last_tool.set(None)
    logger.info(
        "[handler] agent.ainvoke start user=%s session=%s chat_model=%s "
        "timeout=%.0fs llm_timeout=%.0fs msg=%r",
        user_id,
        session_id,
        LLM_MODEL,
        AGENT_INVOKE_TIMEOUT_SEC,
        LLM_REQUEST_TIMEOUT_SEC,
        message_preview[:100],
    )

    task = asyncio.create_task(agent.ainvoke(input_state, config=config))
    t0 = time.monotonic()

    try:
        while True:
            elapsed = time.monotonic() - t0
            remaining = AGENT_INVOKE_TIMEOUT_SEC - elapsed
            if remaining <= 0:
                task.cancel()
                try:
                    await task
                except asyncio.CancelledError:
                    pass
                last_tool = _invoke_last_tool.get()
                logger.error(
                    "[handler] TIMEOUT agent.ainvoke elapsed=%.2fs limit=%.0fs user=%s "
                    "session=%s last_tool=%s chat_model=%s msg=%r",
                    elapsed,
                    AGENT_INVOKE_TIMEOUT_SEC,
                    user_id,
                    session_id,
                    last_tool or "none",
                    LLM_MODEL,
                    message_preview[:100],
                )
                raise asyncio.TimeoutError(
                    f"agent.ainvoke exceeded {AGENT_INVOKE_TIMEOUT_SEC:.0f}s "
                    f"(last_tool={last_tool or 'none'})"
                )

            wait_for = min(AGENT_INVOKE_WARN_INTERVAL_SEC, remaining)
            done, _ = await asyncio.wait({task}, timeout=wait_for)
            if task in done:
                result = task.result()
                elapsed = time.monotonic() - t0
                logger.info(
                    "[handler] agent.ainvoke finished elapsed=%.2fs user=%s session=%s last_tool=%s",
                    elapsed,
                    user_id,
                    session_id,
                    _invoke_last_tool.get() or "none",
                )
                return result

            last_tool = _invoke_last_tool.get()
            logger.warning(
                "[handler] agent.ainvoke slow elapsed=%.2fs user=%s session=%s "
                "last_tool=%s chat_model=%s (likely waiting on LLM after tool)",
                elapsed,
                user_id,
                session_id,
                last_tool or "none",
                LLM_MODEL,
            )
    except Exception:
        if not task.done():
            task.cancel()
            try:
                await task
            except asyncio.CancelledError:
                pass
        raise


@app.entrypoint
async def handler(payload: dict, context: RequestContext) -> dict:
    """Main agent entrypoint with LangChain + Memory support.

    Args:
        payload: JSON body with "message" and optional "file"
        context: Request metadata (session_id, user_id, request_headers)
    """
    # Short-term memory (checkpointer) requires both user_id and session_id
    # to correctly persist and isolate conversation state per user per session.
    if not context.user_id or not context.session_id:
        return {
            "status": "error",
            "error": "Missing required headers: X-GreenNode-AgentBase-User-Id and X-GreenNode-AgentBase-Session-Id are required when using memory.",
        }

    message = payload.get("message", "Hello")
    file_data = payload.get("file")
    t_start = time.monotonic()

    logger.info(
        "[handler] user=%s session=%s msg=%r has_file=%s",
        context.user_id, context.session_id, message[:100], bool(file_data),
    )

    # Handle file upload or external insurance links
    enhanced_message = message
    comparison_question = message.strip() or "So sánh sản phẩm bảo hiểm này với các gói hiện có trên Zalopay"

    if file_data:
        try:
            file_name = file_data.get("name", "unknown")
            logger.info("[handler] processing uploaded file: %s", file_name)

            file_type = file_data.get("type", "")
            raw_content = file_data.get("content")
            cache_key = _build_extraction_cache_key(
                context.user_id,
                context.session_id,
                f"file:{file_name}:{file_data.get('size', 0)}",
            )

            if _is_json_upload(file_type, raw_content):
                json_obj = raw_content if isinstance(raw_content, dict) else json.loads(raw_content)
                insurance_info = resolve_insurance_info(
                    llm=llm_analysis,
                    json_data=json_obj,
                )
            else:
                file_content = parse_uploaded_file(file_data)
                insurance_info = resolve_insurance_info(
                    llm=llm_analysis,
                    text_content=file_content,
                    cache_key=cache_key,
                )

            enhanced_message = build_insurance_comparison_message(
                source_type="file bảo hiểm",
                source_name=file_name,
                insurance_info=insurance_info,
                user_message=comparison_question,
            )

            logger.info("[handler] extracted insurance info: %s", insurance_info.get("product_name", "N/A"))

        except Exception as e:
            logger.error(f"[handler] error processing file: {e}")
            return {
                "status": "error",
                "error": f"Không thể xử lý file upload: {str(e)}",
                "timestamp": datetime.now().isoformat(),
            }
    else:
        urls = extract_urls_from_message(message)
        if urls:
            source_url = urls[0]
            try:
                logger.info("[handler] processing external link: %s", source_url)

                page_content = await fetch_url_content_async(source_url)
                cache_key = _build_extraction_cache_key(
                    context.user_id,
                    context.session_id,
                    f"url:{source_url}",
                )
                insurance_info = resolve_insurance_info(
                    llm=llm_analysis,
                    text_content=page_content,
                    cache_key=cache_key,
                )
                insurance_info["source_url"] = source_url
                enhanced_message = build_insurance_comparison_message(
                    source_type="link bảo hiểm",
                    source_name=source_url,
                    insurance_info=insurance_info,
                    user_message=comparison_question,
                )

                logger.info(
                    "[handler] extracted insurance info from url: %s",
                    insurance_info.get("product_name", "N/A"),
                )
            except Exception as e:
                logger.error(f"[handler] error processing url: {e}")
                return {
                    "status": "error",
                    "error": f"Không thể đọc link ngoài: {str(e)}",
                    "timestamp": datetime.now().isoformat(),
                }

    # Map AgentBase context to LangGraph config
    # thread_id -> session persistence, actor_id -> per-user memory
    config = {
        "configurable": {
            "thread_id": context.session_id,
            "actor_id": context.user_id,
        }
    }

    # ainvoke is required when using async tools (remember, recall)
    try:
        result = await _invoke_agent_with_timeout_logging(
            {"messages": [{"role": "user", "content": enhanced_message}]},
            config,
            user_id=context.user_id,
            session_id=context.session_id,
            message_preview=message,
        )
    except asyncio.TimeoutError:
        elapsed = time.monotonic() - t_start
        logger.error(
            "[handler] TIMEOUT returning error user=%s session=%s elapsed=%.2fs last_tool=%s",
            context.user_id,
            context.session_id,
            elapsed,
            _invoke_last_tool.get() or "none",
        )
        return {
            "status": "error",
            "error": "Agent phản hồi quá chậm. Vui lòng thử lại sau.",
            "timestamp": datetime.now().isoformat(),
        }
    ai_message = result["messages"][-1]

    # Count tool calls made during this invocation for observability
    tool_calls = sum(
        1 for m in result["messages"]
        if hasattr(m, "type") and m.type == "tool"
    )
    elapsed = time.monotonic() - t_start
    logger.info(
        "[handler] done user=%s tool_calls=%d elapsed=%.2fs has_file=%s has_url=%s",
        context.user_id,
        tool_calls,
        elapsed,
        bool(file_data),
        bool(not file_data and extract_urls_from_message(message)),
    )

    return {
        "status": "success",
        "response": normalize_brand_name(ai_message.content),
        "timestamp": datetime.now().isoformat(),
    }


@app.ping
def health_check() -> PingStatus:
    """Custom health check for GET /health endpoint."""
    return PingStatus.HEALTHY


if __name__ == "__main__":
    port = int(os.environ.get("AGENT_PORT", 8080))
    app.run(port=port, host="0.0.0.0")
